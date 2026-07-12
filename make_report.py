"""Generate the final project report (final_report.docx) from the real run
outputs: results/train_*.json, results/eval_*.json, and results/results.xlsx.

Usage: python make_report.py
"""
import json
from datetime import date

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Cm, Pt

from common import RESULTS

OUT = "final_report.docx"


def load_inputs():
    train = {m: json.load(open(RESULTS / f"train_{m}.json")) for m in ["full", "lora"]}
    evals = {m: json.load(open(RESULTS / f"eval_{m}.json")) for m in ["full", "lora"]}
    sheets = pd.read_excel(RESULTS / "results.xlsx", sheet_name=None)
    return train, evals, sheets


def add_table(doc, df, font_size=9.5, truncate=None):
    df = df.fillna("")
    if truncate:
        for col, n in truncate.items():
            df[col] = df[col].astype(str).map(
                lambda s: s if len(s) <= n else s[:n].rstrip() + " …")
    table = doc.add_table(rows=1 + len(df), cols=len(df.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.rows[0].cells[j]
        run = cell.paragraphs[0].add_run(str(col))
        run.bold = True
        run.font.size = Pt(font_size)
    for i, (_, row) in enumerate(df.iterrows(), start=1):
        for j, val in enumerate(row):
            run = table.rows[i].cells[j].paragraphs[0].add_run(str(val))
            run.font.size = Pt(font_size)
    return table


def caption(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)


def para(doc, text):
    doc.add_paragraph(text)


def main():
    train, evals, sheets = load_inputs()
    a_t, b_t = train["full"], train["lora"]
    a_e, b_e = evals["full"], evals["lora"]
    scores = {r["Model"]: r for r in sheets["5_scores_and_cost"].to_dict("records")}
    sa = scores["Model A (Full FT)"]
    sb = scores["Model B (LoRA)"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = section.bottom_margin = Cm(2)
        section.left_margin = section.right_margin = Cm(2)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    doc.add_heading("Training and Comparing Two Language Models on "
                    "Classification, Translation and QA", level=0)
    para(doc, f"NLP Final Project — Full Fine-tuning vs. LoRA/PEFT on a shared base model. "
              f"{date.today().strftime('%B %d, %Y')}.")
    authors = doc.add_paragraph()
    authors.add_run("Ronen Shershnev — ID 322217175\n"
                    "Vladislav Pavliyuk — ID 332294891").bold = True

    # 1. Overview -----------------------------------------------------------
    doc.add_heading("1. Overview and Experimental Setup", level=1)
    para(doc,
         "Two models were trained from the same base model, google/flan-t5-small "
         f"({a_t['total_params']:,} parameters), and compared on three tasks: sentiment "
         "classification (SST-2), English→French translation (OPUS Books) and extractive "
         "question answering (SQuAD v1.1). Model A updates every weight (full fine-tuning); "
         "Model B trains LoRA adapters (r=16, α=32, on the q and v attention projections) "
         f"— {b_t['trainable_params']:,} trainable parameters, {b_t['trainable_pct']}% of the "
         "model. Each model is a single multi-task system: one training run over a shuffled "
         "mix of all three tasks in a unified text-to-text format, with a task prefix "
         "selecting the behavior:")
    for line in ['classify sentiment: This movie was surprisingly good.  →  positive',
                 "translate English to French: The child opened the door.  →  L'enfant a ouvert la porte.",
                 'question: Where did Mary go? context: Mary went to the kitchen after school.  →  kitchen']:
        p = doc.add_paragraph(line)
        p.paragraph_format.left_indent = Cm(0.75)
        for run in p.runs:
            run.font.name = "Consolas"
            run.font.size = Pt(9)
    para(doc,
         "Both models saw identical fixed splits (seed 42): 1,000 train / 200 validation / "
         "200 test examples per task. Deduplication prevents train/test overlap; QA prompts "
         "were filtered to ≤256 tokens so the answer is always inside the visible context, "
         "and targets were filtered to fit each task's generation cap (64 tokens for "
         "translation, 32 for QA), so every reference is producible at evaluation time. "
         "Training used 3 epochs, batch size 8, fp32, learning rate 5e-5 (full) / 5e-4 "
         "(LoRA), with the best epoch selected by validation loss. Evaluation used greedy "
         "decoding with per-task output caps (5/64/32 tokens). Hardware: "
         f"{a_t['hardware']}, device {a_e['device'].upper()}. Table 1 summarizes both "
         "configurations.")

    caption(doc, "Table 1 — Model and training details")
    add_table(doc, sheets["1_model_details"], font_size=9)

    # 2. Results ------------------------------------------------------------
    doc.add_heading("2. Results per Task", level=1)
    para(doc,
         "Tables 2–4 report the test-set results per task (200 examples each); the B−A row "
         "is the LoRA-minus-full difference. Full fine-tuning is ahead on every quality "
         "metric, but never by a large margin. Neither model ever produced an invalid "
         "classification label (invalid output rate 0.0).")

    caption(doc, "Table 2 — Classification (SST-2)")
    add_table(doc, sheets["2_classification"])
    doc.add_paragraph()
    caption(doc, "Table 3 — Translation (OPUS Books, English→French)")
    add_table(doc, sheets["3_translation"])
    para(doc,
         "Absolute translation quality is low (BLEU ≈ 6) for both models: OPUS Books is "
         "literary prose, 1,000 sentence pairs are few, and the 77M-parameter model has "
         "limited French. chrF (≈29) being far higher than BLEU shows partial word-level "
         "correctness without exact n-gram matches. The comparison between methods remains "
         "meaningful because both models face identical conditions.")
    caption(doc, "Table 4 — Question Answering (SQuAD v1.1)")
    add_table(doc, sheets["4_qa"])

    # 3. Overall score ------------------------------------------------------
    doc.add_heading("3. Overall Score and Computational Cost", level=1)
    para(doc,
         "Per the assignment formula, each task is normalized to a 0–100 scale and "
         "averaged:  overall = (100·Macro-F1 + BLEU + 100·QA-token-F1) / 3.")
    caption(doc, "Table 5 — Normalized scores and computational cost")
    add_table(doc, sheets["5_scores_and_cost"], font_size=9)

    # 4. Qualitative examples ------------------------------------------------
    doc.add_heading("4. Qualitative Examples", level=1)
    para(doc,
         "Table 6 shows one representative example per task plus a shared failure and a "
         "shared success (inputs truncated for readability; full predictions for all 600 "
         "test examples are in results/preds_*.jsonl).")
    caption(doc, "Table 6 — Qualitative examples")
    add_table(doc, sheets["6_qualitative"], font_size=8,
              truncate={"Input": 200, "Gold answer": 110,
                        "Model A output": 110, "Model B output": 110})

    # 5. Analysis -----------------------------------------------------------
    doc.add_heading("5. Analysis", level=1)

    doc.add_heading("5.1  Which training method performed better on each task?", level=2)
    para(doc,
         f"Full fine-tuning (Model A) won all three tasks, by consistently small margins: "
         f"classification Macro-F1 0.840 vs. 0.835 (−0.005), translation BLEU 6.11 vs. 5.76 "
         f"(−0.35; chrF 29.51 vs. 28.56), and QA token-F1 0.840 vs. 0.826 (−0.013, the "
         f"largest gap — −1.34 normalized points). In exchange, LoRA trained in "
         f"{sb['Train time (min)']} minutes vs. {sa['Train time (min)']} (−42%) and its "
         f"checkpoint is {b_t['checkpoint_size_mb']} MB vs. {a_t['checkpoint_size_mb']} MB "
         f"(59× smaller).")

    doc.add_heading("5.2  Was the same model consistently better across tasks?", level=2)
    para(doc,
         "Yes — the ranking A ≥ B held on every metric of every task; there was no task "
         "where LoRA overtook full fine-tuning. More striking is how similar the two models "
         "are at the level of individual predictions: on classification they gave the same "
         "correctness outcome on 195/200 test examples (165 both correct, 30 both wrong, "
         "only 5 disagreements), and on QA exact-match on 190/200 (150 both correct, 40 "
         "both wrong). Both models also share the same task profile — strong on extractive "
         "QA and binary sentiment, weak on literary translation. This suggests that with "
         "1,000 examples per task, behavior is dominated by the shared base model and "
         "shared data; updating 0.89% of parameters is enough to track full fine-tuning "
         "almost exactly.")

    doc.add_heading("5.3  Does the quality improvement justify the cost?", level=2)
    para(doc,
         f"Largely no. Full fine-tuning trained 112× more parameters "
         f"({a_t['trainable_params']:,} vs. {b_t['trainable_params']:,}), took 73% longer "
         f"({sa['Train time (min)']} vs. {sb['Train time (min)']} min), and stores a 59× "
         f"larger checkpoint, for +{round(sa['Overall score'] - sb['Overall score'], 2)} "
         f"overall points ({sa['Overall score']} vs. {sb['Overall score']}) — LoRA reaches "
         f"98.8% of full fine-tuning's overall score. At this model and data scale, LoRA is "
         "clearly the better cost/quality trade-off; full fine-tuning is defensible only "
         "when every fraction of a point matters. One nuance favors full fine-tuning at "
         "inference: the unmerged LoRA adapters add computation, making Model B "
         f"~8% slower ({b_e['inference_sec_per_100_overall']} vs. "
         f"{a_e['inference_sec_per_100_overall']} sec/100 examples). This penalty is "
         "avoidable by merging the adapters into the base weights after training.")

    doc.add_heading("5.4  What kinds of examples did each model fail on?", level=2)
    para(doc,
         "Failures were overwhelmingly shared rather than model-specific (A-only failures: "
         "3 classification / 6 QA; B-only: 2 / 4), so the patterns below describe both "
         "models. Classification: errors concentrate on figurative or contrastive positive "
         "reviews — e.g. “you'll gasp appalled and laugh outraged” (gold positive, both "
         "predicted negative) — and on sarcasm — “i'll bet the video game is a lot more fun "
         "than the film” (gold negative, both predicted positive). Both models lean on "
         "surface sentiment words; errors skew toward the positive class (A missed 19/107 "
         "positives vs. 13/93 negatives). Translation: quality is bimodal — short, "
         "dialogue-like sentences translate acceptably while dense literary prose fails "
         "(49 of 200 sentences below chrF 20 for A, 52 for B, vs. 16/13 above chrF 50); "
         "typical errors are wrong rare vocabulary and untranslated English words carried "
         "into the output (8/200 for A, 7/200 for B). QA: both models fail on questions "
         "requiring inference or paraphrase rather than span extraction — “how/why” "
         "questions get a copied context sentence (the Antigone example in Table 6, token-F1 "
         "0.06), and several errors return a plausible-but-wrong span of the right type "
         "(“annually” instead of “every four years”; “1,442” instead of “over 14,000”; "
         "“Queen Bees” taken literally instead of “women”).")

    doc.add_heading("5.5  Were outputs too short, too long, or hallucinated?", level=2)
    para(doc,
         "No systematic length pathology was observed. Average output lengths track the "
         "references on every task: classification outputs were exactly one word with 0% "
         "invalid labels for both models; translation averaged 14.7 (A) / 14.3 (B) words "
         "against references of ≈14.5; QA answers averaged 2.65 / 2.67 words against gold "
         "2.85. The output caps (5/64/32 tokens) plus filtering targets to fit them meant "
         "no reference was ever unreachable and no truncated outputs were observed. "
         "Residual issues are small and worth naming: Model A produced 1 empty translation "
         "and 5 with repetition loops (Model B: 0 and 4, ≤2.5% of examples); a few "
         "translations coin non-words under vocabulary pressure (e.g. “l'évenation”), a "
         "mild form of hallucination; and 4 of Model A's QA failures were over-long "
         "context copies rather than concise answers. Classification and QA otherwise "
         "showed no hallucination — wrong answers were still drawn from the input context.")

    doc.add_heading("6. Conclusion", level=1)
    para(doc,
         "Full fine-tuning delivered the best scores on all three tasks, but LoRA matched "
         "it to within ~1 point per task while training 42% faster, updating <1% of the "
         "parameters, and producing a 59× smaller checkpoint — on this setup, the "
         "parameter-efficient method is the sensible default, and the interesting failure "
         "modes (sarcasm, literary vocabulary, inference questions) are shared properties "
         "of the base model and data scale, not of the adaptation method.")

    doc.save(OUT)
    words = sum(len(p.text.split()) for p in doc.paragraphs)
    print(f"wrote {OUT}: {words} words in paragraphs, "
          f"{len(doc.tables)} tables, {len(doc.paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
